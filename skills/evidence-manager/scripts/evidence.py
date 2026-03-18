#!/usr/bin/env python3
"""
Evidence Manager: Master evidence index for legal proceedings.
Tracks exhibits across bar complaints, immigration cases, and appeals.
"""

import os
import sys
import json
import hashlib
import argparse
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from typing import Dict, List, Set, Tuple, Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configuration
SUPPORTED_EVIDENCE_TYPES = {
    'pdf': {'ext': ['.pdf']},
    'docx': {'ext': ['.docx', '.doc']},
    'xlsx': {'ext': ['.xlsx', '.xls']},
    'txt': {'ext': ['.txt']},
    'rtf': {'ext': ['.rtf']},
    'image': {'ext': ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.gif', '.bmp']},
    'audio': {'ext': ['.mp3', '.wav', '.aac', '.flac', '.m4a']},
    'video': {'ext': ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv']},
    'zip': {'ext': ['.zip', '.rar', '.7z']},
}

KNOWN_FOLDERS = {
    'shared_evidence': 'BAR COMPLAINTS/10.SUPPORTING EVIDENCE',
    'bar_complaints': 'BAR COMPLAINTS',
    'immigration': 'CASES EB1 AND EB2',
    'binders': 'Binders',
    'camargo': 'camargo-reply',
    'aggarwal': 'julia_response_malika',
}


class EvidenceManager:
    def __init__(self, workspace_root: str):
        self.workspace_root = Path(workspace_root).resolve()
        self.index_path = self.workspace_root / 'evidence_index.json'
        self.index = self._load_index()

    def _load_index(self) -> Dict:
        """Load evidence index from JSON, or create empty structure."""
        if self.index_path.exists():
            try:
                with open(self.index_path, 'r') as f:
                    return json.load(f)
            except (json.JSONDecodeError, IOError):
                print(f"Warning: Could not load {self.index_path}, starting fresh")
                return {'exhibits': {}, 'metadata': {'last_scanned': None}}
        return {'exhibits': {}, 'metadata': {'last_scanned': None}}

    def _save_index(self):
        """Save evidence index to JSON."""
        self.index['metadata']['last_scanned'] = datetime.now().isoformat()
        with open(self.index_path, 'w') as f:
            json.dump(self.index, f, indent=2)
        print(f"Index saved to {self.index_path}")

    def _get_file_type(self, filepath: Path) -> str:
        """Determine file type based on extension."""
        ext = filepath.suffix.lower()
        for ftype, config in SUPPORTED_EVIDENCE_TYPES.items():
            if ext in config['ext']:
                return ftype
        return 'unknown'

    def _compute_hash(self, filepath: Path) -> str:
        """Compute SHA256 hash of file for duplicate detection."""
        sha256_hash = hashlib.sha256()
        try:
            with open(filepath, 'rb') as f:
                for byte_block in iter(lambda: f.read(4096), b''):
                    sha256_hash.update(byte_block)
            return sha256_hash.hexdigest()
        except (IOError, OSError) as e:
            print(f"Warning: Could not hash {filepath}: {e}")
            return ''

    def scan(self, verbose: bool = False):
        """
        Recursively scan workspace for evidence files.
        Auto-discover and index all supported file types.
        """
        print(f"Scanning {self.workspace_root} for evidence files...")
        found_count = 0
        updated_count = 0

        for filepath in self.workspace_root.rglob('*'):
            if filepath.is_file() and self._is_evidence_file(filepath):
                found_count += 1
                file_type = self._get_file_type(filepath)
                rel_path = filepath.relative_to(self.workspace_root).as_posix()
                file_hash = self._compute_hash(filepath)

                exhibit_id = self._generate_exhibit_id()

                # Check if file already indexed (by hash or path)
                existing_id = self._find_indexed_file(filepath, file_hash)
                if existing_id:
                    if verbose:
                        print(f"  [EXISTING] {rel_path}")
                    continue

                # Create new index entry
                entry = {
                    'id': exhibit_id,
                    'filename': filepath.name,
                    'description': f"[AUTO] {filepath.name}",
                    'file_path': rel_path,
                    'date_created': datetime.fromtimestamp(filepath.stat().st_mtime).isoformat()[:10],
                    'file_type': file_type,
                    'file_hash': file_hash,
                    'proceedings': self._infer_proceedings(rel_path),
                    'claims': [],
                    'tags': [],
                    'date_indexed': datetime.now().isoformat()[:10],
                    'notes': '',
                }

                self.index['exhibits'][exhibit_id] = entry
                updated_count += 1

                if verbose:
                    print(f"  [NEW] {rel_path} -> {exhibit_id}")

        print(f"Scan complete: {found_count} evidence files found, {updated_count} new entries added")
        self._save_index()

    def _is_evidence_file(self, filepath: Path) -> bool:
        """Check if file is a supported evidence type."""
        if filepath.name.startswith('.'):
            return False
        for config in SUPPORTED_EVIDENCE_TYPES.values():
            if filepath.suffix.lower() in config['ext']:
                return True
        return False

    def _generate_exhibit_id(self) -> str:
        """Generate next available exhibit ID."""
        existing_ids = [e.get('id', '').replace('EX-', '') for e in self.index['exhibits'].values()]
        max_num = 0
        for id_str in existing_ids:
            try:
                max_num = max(max_num, int(id_str))
            except ValueError:
                continue
        return f"EX-{str(max_num + 1).zfill(3)}"

    def _find_indexed_file(self, filepath: Path, file_hash: str) -> Optional[str]:
        """Find if file is already indexed by EXACT PATH (not hash).

        We only skip files with the same path (already indexed).
        Files with the same hash but different paths are BOTH indexed —
        this is essential for duplicate detection to work correctly.
        """
        rel_path = filepath.relative_to(self.workspace_root).as_posix()
        for exhibit_id, entry in self.index['exhibits'].items():
            if entry['file_path'] == rel_path:
                return exhibit_id
        return None

    def _infer_proceedings(self, rel_path: str) -> List[str]:
        """Infer which proceedings a file relates to based on folder path."""
        proceedings = []
        path_lower = rel_path.lower()

        if 'attorney_smith' in path_lower or 'smith' in path_lower:
            proceedings.append('Smith_Bar_Complaint')
        if 'camargo' in path_lower:
            proceedings.append('Camargo_Proceeding')
        if 'aggarwal' in path_lower or 'malika' in path_lower:
            proceedings.append('Aggarwal_Bar_Complaint')
        if 'widmer' in path_lower or 'jasmine' in path_lower:
            proceedings.append('Widmer_NJ_Bar')
        if 'wasung' in path_lower or 'pieter' in path_lower:
            proceedings.append('Wasung_DC_Bar')
        if 'rutahweire' in path_lower or 'phillip' in path_lower or 'philip' in path_lower:
            proceedings.append('Rutahweire_MA_Bar')
        if 'hayman' in path_lower:
            proceedings.append('Hayman_Woodward_General')
        if 'eb1' in path_lower or 'cases eb' in path_lower:
            proceedings.append('EB1A_Immigration')
        if 'eb2' in path_lower or 'cases eb' in path_lower:
            proceedings.append('EB2_NIW_Immigration')
        if 'i-485' in path_lower or 'i485' in path_lower:
            proceedings.append('I485_Adjustment')
        if 'i-290' in path_lower or 'i290' in path_lower or 'appeal' in path_lower:
            proceedings.append('I290B_Appeal')
        if 'foia' in path_lower:
            proceedings.append('FOIA_Request')
        if 'binders' in path_lower:
            proceedings.append('Appeal_Proceedings')

        return proceedings

    def list_exhibits(self, proceeding: Optional[str] = None,
                     claim: Optional[str] = None,
                     file_type: Optional[str] = None,
                     date_from: Optional[str] = None,
                     date_to: Optional[str] = None) -> List[Dict]:
        """List exhibits with optional filters."""
        results = []
        for exhibit_id, entry in self.index['exhibits'].items():
            if proceeding and proceeding not in entry.get('proceedings', []):
                continue
            if claim and claim not in entry.get('claims', []):
                continue
            if file_type and entry.get('file_type') != file_type:
                continue
            if date_from and entry.get('date_created', '') < date_from:
                continue
            if date_to and entry.get('date_created', '') > date_to:
                continue
            results.append(entry)
        return sorted(results, key=lambda x: x.get('id', ''))

    def search(self, query: str) -> List[Dict]:
        """Search exhibits by keyword in description or filename."""
        query_lower = query.lower()
        results = []
        for exhibit_id, entry in self.index['exhibits'].items():
            if (query_lower in entry.get('description', '').lower() or
                query_lower in entry.get('filename', '').lower() or
                query_lower in ' '.join(entry.get('claims', [])).lower()):
                results.append(entry)
        return results

    def add_exhibit(self, filepath: str, description: str,
                   proceedings: Optional[List[str]] = None,
                   claims: Optional[List[str]] = None):
        """Manually add or update an exhibit."""
        full_path = self.workspace_root / filepath
        if not full_path.exists():
            print(f"Error: File not found: {full_path}")
            return False

        exhibit_id = self._find_indexed_file(full_path, '')
        if not exhibit_id:
            exhibit_id = self._generate_exhibit_id()

        rel_path = full_path.relative_to(self.workspace_root).as_posix()
        file_hash = self._compute_hash(full_path)

        entry = {
            'id': exhibit_id,
            'filename': full_path.name,
            'description': description,
            'file_path': rel_path,
            'date_created': datetime.fromtimestamp(full_path.stat().st_mtime).isoformat()[:10],
            'file_type': self._get_file_type(full_path),
            'file_hash': file_hash,
            'proceedings': proceedings or self._infer_proceedings(rel_path),
            'claims': claims or [],
            'tags': [],
            'date_indexed': datetime.now().isoformat()[:10],
            'notes': '',
        }

        self.index['exhibits'][exhibit_id] = entry
        self._save_index()
        print(f"Added/updated exhibit {exhibit_id}: {description}")
        return True

    def tag_exhibit(self, exhibit_id: str,
                   proceedings: Optional[List[str]] = None,
                   claims: Optional[List[str]] = None,
                   tags: Optional[List[str]] = None,
                   notes: Optional[str] = None):
        """Update tags, proceedings, and claims for an exhibit."""
        if exhibit_id not in self.index['exhibits']:
            print(f"Error: Exhibit {exhibit_id} not found")
            return False

        entry = self.index['exhibits'][exhibit_id]
        if proceedings:
            entry['proceedings'] = list(set(entry.get('proceedings', []) + proceedings))
        if claims:
            entry['claims'] = list(set(entry.get('claims', []) + claims))
        if tags:
            entry['tags'] = list(set(entry.get('tags', []) + tags))
        if notes:
            entry['notes'] = notes

        self._save_index()
        print(f"Tagged exhibit {exhibit_id}")
        return True

    def cross_reference(self, proceeding: Optional[str] = None,
                       claim: Optional[str] = None) -> Dict:
        """Show cross-references for a proceeding or claim."""
        if proceeding:
            exhibits = self.list_exhibits(proceeding=proceeding)
            print(f"\n=== Evidence for {proceeding} ===")
            for entry in exhibits:
                print(f"\n{entry['id']}: {entry['description']}")
                print(f"  File: {entry['file_path']}")
                print(f"  Claims: {', '.join(entry['claims']) or 'None specified'}")
            return {'proceeding': proceeding, 'count': len(exhibits), 'exhibits': exhibits}

        if claim:
            exhibits = self.list_exhibits(claim=claim)
            print(f"\n=== Evidence supporting claim: {claim} ===")
            for entry in exhibits:
                print(f"\n{entry['id']}: {entry['description']}")
                print(f"  Proceedings: {', '.join(entry['proceedings'])}")
                print(f"  File: {entry['file_path']}")
            return {'claim': claim, 'count': len(exhibits), 'exhibits': exhibits}

        return {}

    def find_duplicates(self) -> Dict[str, List[str]]:
        """Identify duplicate files by hash."""
        hash_map = defaultdict(list)
        for exhibit_id, entry in self.index['exhibits'].items():
            file_hash = entry.get('file_hash', '')
            if file_hash:
                hash_map[file_hash].append({
                    'id': exhibit_id,
                    'path': entry['file_path'],
                    'filename': entry['filename']
                })

        duplicates = {h: files for h, files in hash_map.items() if len(files) > 1}

        if duplicates:
            print("\n=== DUPLICATE FILES FOUND ===")
            for file_hash, files in duplicates.items():
                print(f"\n[Hash: {file_hash[:8]}...]")
                for f in files:
                    print(f"  {f['id']}: {f['path']}")
                print(f"  Action: Keep primary copy, consolidate or remove duplicates")
        else:
            print("No duplicates found.")

        return duplicates

    def find_gaps(self, proceeding: str) -> Dict:
        """Analyze claims lacking supporting evidence in a proceeding."""
        exhibits = self.list_exhibits(proceeding=proceeding)

        # Collect all claims mentioned in exhibits for this proceeding
        all_claims = set()
        claims_with_evidence = defaultdict(list)
        for entry in exhibits:
            for claim in entry.get('claims', []):
                all_claims.add(claim)
                claims_with_evidence[claim].append(entry['id'])

        gaps = {}
        print(f"\n=== Gap Analysis for {proceeding} ===")

        if not all_claims:
            print("No claims found for this proceeding. Add claims to exhibits via tag command.")
            return gaps

        print(f"\nAnalyzed {len(exhibits)} exhibits covering {len(all_claims)} distinct claims:\n")

        for claim in sorted(all_claims):
            evidence_ids = claims_with_evidence.get(claim, [])
            if not evidence_ids:
                gaps[claim] = {
                    'status': 'CRITICAL GAP',
                    'count': 0,
                    'exhibits': []
                }
                print(f"  ⚠ {claim}: [CRITICAL] 0 exhibits")
            elif len(evidence_ids) == 1:
                gaps[claim] = {
                    'status': 'THIN',
                    'count': 1,
                    'exhibits': evidence_ids
                }
                print(f"  ⚠ {claim}: THIN ({len(evidence_ids)} exhibit: {', '.join(evidence_ids)})")
            else:
                gaps[claim] = {
                    'status': 'ADEQUATE',
                    'count': len(evidence_ids),
                    'exhibits': evidence_ids
                }
                print(f"  ✓ {claim}: {len(evidence_ids)} exhibits")

        critical_gaps = {c: g for c, g in gaps.items() if g['status'] == 'CRITICAL GAP'}
        if critical_gaps:
            print(f"\nACTION REQUIRED: {len(critical_gaps)} claim(s) lack supporting evidence:")
            for claim in critical_gaps:
                print(f"  - {claim}")

        return gaps

    def generate_docx_index(self, proceeding: str, output_path: Optional[str] = None) -> str:
        """Generate formatted evidence index as a .docx file."""
        if not DOCX_AVAILABLE:
            print("Error: python-docx not installed. Install with: pip install python-docx")
            return ''

        exhibits = self.list_exhibits(proceeding=proceeding)
        if not exhibits:
            print(f"No exhibits found for {proceeding}")
            return ''

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading(f'Evidence Index: {proceeding}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        meta.add_run(f'Total Exhibits: {len(exhibits)}\n')
        meta.add_run(f'Proceeding: {proceeding}')

        doc.add_paragraph()  # Spacer

        # Table of contents / summary by claim
        claims_map = defaultdict(list)
        for entry in exhibits:
            for claim in entry.get('claims', []):
                claims_map[claim].append(entry['id'])

        if claims_map:
            doc.add_heading('Claims Summary', level=2)
            for claim in sorted(claims_map.keys()):
                p = doc.add_paragraph(f'{claim}', style='List Bullet')
                p.add_run(f' ({len(claims_map[claim])} exhibits): ').italic = True
                p.add_run(', '.join(claims_map[claim]))

        doc.add_page_break()

        # Exhibit listing
        doc.add_heading('Exhibits', level=2)

        for idx, entry in enumerate(exhibits, 1):
            # Exhibit header - generate letter ID (A-Z, AA-AZ, BA-BZ, etc.)
            def get_letter_id(num):
                """Convert number to letter ID: 1->A, 26->Z, 27->AA, etc."""
                result = ""
                num = num - 1  # 0-indexed
                while num >= 0:
                    result = chr(65 + (num % 26)) + result
                    num = num // 26 - 1
                    if num < 0:
                        break
                return result

            heading = doc.add_heading(f"Exhibit {get_letter_id(idx)}: {entry['id']}", level=3)

            # Details
            details = doc.add_paragraph()
            details.add_run('Description: ').bold = True
            details.add_run(entry.get('description', 'N/A'))

            details = doc.add_paragraph()
            details.add_run('File: ').bold = True
            details.add_run(entry.get('file_path', 'N/A'))

            details = doc.add_paragraph()
            details.add_run('Type: ').bold = True
            details.add_run(entry.get('file_type', 'unknown').upper())

            if entry.get('date_created'):
                details = doc.add_paragraph()
                details.add_run('Date: ').bold = True
                details.add_run(entry['date_created'])

            if entry.get('claims'):
                details = doc.add_paragraph()
                details.add_run('Supports Claims: ').bold = True
                details.add_run(', '.join(entry['claims']))

            if entry.get('notes'):
                details = doc.add_paragraph()
                details.add_run('Notes: ').bold = True
                details.add_run(entry['notes'])

            doc.add_paragraph()  # Spacer

        # Save
        if not output_path:
            output_path = str(self.workspace_root / f'{proceeding}_Evidence_Index.docx')

        doc.save(output_path)
        print(f"\nGenerated evidence index: {output_path}")
        return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Evidence Manager: Manage exhibits and evidence across legal proceedings'
    )
    parser.add_argument('--workspace', default=os.getcwd(),
                       help='Workspace root directory (default: current directory)')

    subparsers = parser.add_subparsers(dest='command', help='Commands')

    # scan
    subparsers.add_parser('scan', help='Scan workspace for evidence files')

    # list
    list_parser = subparsers.add_parser('list', help='List indexed exhibits')
    list_parser.add_argument('--proceeding', help='Filter by proceeding')
    list_parser.add_argument('--claim', help='Filter by claim')
    list_parser.add_argument('--type', help='Filter by file type')
    list_parser.add_argument('--from', dest='date_from', help='Filter by date (YYYY-MM-DD)')
    list_parser.add_argument('--to', dest='date_to', help='Filter by date (YYYY-MM-DD)')

    # search
    search_parser = subparsers.add_parser('search', help='Search exhibits by keyword')
    search_parser.add_argument('query', help='Search query')

    # add
    add_parser = subparsers.add_parser('add', help='Manually add/update an exhibit')
    add_parser.add_argument('filepath', help='Path to exhibit file')
    add_parser.add_argument('description', help='Description of exhibit')
    add_parser.add_argument('--proceeding', action='append', help='Proceeding (repeatable)')
    add_parser.add_argument('--claim', action='append', help='Claim (repeatable)')

    # tag
    tag_parser = subparsers.add_parser('tag', help='Update tags/claims for an exhibit')
    tag_parser.add_argument('exhibit_id', help='Exhibit ID (e.g., EX-001)')
    tag_parser.add_argument('--proceeding', action='append', help='Add proceeding (repeatable)')
    tag_parser.add_argument('--claim', action='append', help='Add claim (repeatable)')
    tag_parser.add_argument('--tag', action='append', help='Add tag (repeatable)')
    tag_parser.add_argument('--notes', help='Add notes')

    # cross-reference
    xref_parser = subparsers.add_parser('cross-reference', help='Show cross-references')
    xref_parser.add_argument('--proceeding', help='Filter by proceeding')
    xref_parser.add_argument('--claim', help='Filter by claim')

    # duplicates
    subparsers.add_parser('duplicates', help='Find duplicate files')

    # gaps
    gaps_parser = subparsers.add_parser('gaps', help='Analyze claims lacking evidence')
    gaps_parser.add_argument('proceeding', help='Proceeding name')

    # generate-index
    gen_parser = subparsers.add_parser('generate-index', help='Generate formatted evidence index')
    gen_parser.add_argument('proceeding', help='Proceeding name')
    gen_parser.add_argument('--output', help='Output file path')

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return 1

    manager = EvidenceManager(args.workspace)

    try:
        if args.command == 'scan':
            manager.scan(verbose=False)

        elif args.command == 'list':
            exhibits = manager.list_exhibits(
                proceeding=args.proceeding,
                claim=args.claim,
                file_type=getattr(args, 'type', None),
                date_from=getattr(args, 'date_from', None),
                date_to=getattr(args, 'date_to', None)
            )
            print(f"\n=== Indexed Exhibits ({len(exhibits)}) ===\n")
            for entry in exhibits:
                print(f"{entry['id']}: {entry['description']}")
                print(f"  Path: {entry['file_path']}")
                print(f"  Type: {entry['file_type']}")
                if entry['proceedings']:
                    print(f"  Proceedings: {', '.join(entry['proceedings'])}")
                if entry['claims']:
                    print(f"  Claims: {', '.join(entry['claims'])}")
                print()

        elif args.command == 'search':
            results = manager.search(args.query)
            print(f"\n=== Search Results for '{args.query}' ({len(results)}) ===\n")
            for entry in results:
                print(f"{entry['id']}: {entry['description']}")
                print(f"  Path: {entry['file_path']}")
                print()

        elif args.command == 'add':
            manager.add_exhibit(
                args.filepath,
                args.description,
                proceedings=args.proceeding,
                claims=args.claim
            )

        elif args.command == 'tag':
            manager.tag_exhibit(
                args.exhibit_id,
                proceedings=args.proceeding,
                claims=args.claim,
                tags=args.tag,
                notes=args.notes
            )

        elif args.command == 'cross-reference':
            manager.cross_reference(
                proceeding=args.proceeding,
                claim=args.claim
            )

        elif args.command == 'duplicates':
            manager.find_duplicates()

        elif args.command == 'gaps':
            manager.find_gaps(args.proceeding)

        elif args.command == 'generate-index':
            manager.generate_docx_index(args.proceeding, output_path=args.output)

        return 0

    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
